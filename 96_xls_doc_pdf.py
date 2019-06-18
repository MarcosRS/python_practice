# Workbook
# .xls extension
# Sheets/Worksheets
# Columns & Rows
# Cell

# pip install openpyxl (thirdparty)

import openpyxl
import os

os.chdir('./somefolder/')
workbook = openpyxl.load_workbook('example.xlsx')
type(workbook) # to check if it's a workbook

sheet = worksheet.get_sheet_by_name('Sheet1')
type(workbook) # to check if it's sheet

 namesOfSheetsList = worksheet.get_sheet_by_names()

 cellA1 = sheet['A1'] # gets a cell object

cellA1.value # automatically transforms. example '4/5/2014 13:34' =  datetime.datetime(2014,4,5,13,34,2)
str(cellA1.value) # 4/5/2014 13:34 returns exact text from cell

# you can also use : (to get cells)
cellA1_2 = sheet.cell(row = 1, column = 1)
#cellA1_2 == cellA1



# EDITING SPREADSHEET
import openpyxl , os
wb = openpyxl.Workbook() # creates a new one
wb.get_sheet_names() # ['Sheet']
sheet = wb.get_sheet_name('Sheet')
sheet['A1'] == None  # True

#to save info (just in memory for now)
sheet['A1'] = 42
sheet['A1'] = 'Hello'

os.chdir('./folder') #specify path
wb.save('example.xlsx') #save it with a different filename if you had a previous
# to avoid overwritting

# Add more sheets

sheet2 = wb.create_sheet() 
sheet2.title # Sheet2 (defautl name)
sheet2.title = 'mysheet2'
wb.save('example2.xlsx') # assign name

sheet3 = wb.create_sheet(index = 0, title = 'mythirdsheet') # this creates a new sheet and places it as the first (position) sheet
wb.save('example3.xlsx')





#PDF - Reader
import PyPDF2 # third party , the best one
import os

os.chdir('./users/al/documents')

pdfFile = open('somepdf.pdf', 'rb') # read binary mode

reader =  PyPDF2.PdFileReader(pdfFile)
reader.numPages

page = reader.getPage(0)
page.extractText() # sometime you have to playaround with the deta returned since it not 100% exact

#to read all pages

for pageNum in range(reader.numPages):
    print(reader.getPage(pageNum).extractText())


#PDF - Writing 
# Very limited - only edits at page level, placement

pdfFile1 = open('somepdf1.pdf', 'rb') 
pdfFile2 = open('somepdf2.pdf', 'rb') 

reader1 =  PyPDF2.PdFileReader(pdfFile1)
reader2 =  PyPDF2.PdFileReader(pdfFile2)

writer =  PyPDF2.PdFileWriter()

for pageNum in range(reader1.numPages):
    page = reader1.getPage(pageNum)
    writer.addPage(page)
    
for pageNum in range(reader2.numPages):
    page = reader2.getPage(pageNum)
    writer.addPage(page)



outputFile = open('combinedpdf.pdf', 'wb') #write binary mode
write.write(outputFile)

pdfFile1.close()
pdfFile2.close()





#WORD DOC 
import docx # pip install python-docx 
d = docx.Documents('/document.docx')


d.paragraph[0] # its an object
d.paragraph[0].text

p = d.paragraph[1]
p.runs # its whenever style changes within the paragraph

p.runs[0].text # the are placed in a list
p.runs[0].bold # True or False 

# you can change 
p.runs[0].text = '123'
p.runs[0].underline = True
p.style #'Normal'

d.save('/document2.docx')


doc = docx.Documents() #creates new document
doc.add_paragraph('Imagine this is a big paragraph')
doc.add_paragraph('Imagine this is a big paragraph 2')
doc.save('/document3.docx')
p = doc.paragraph[0]
p.add_run('This is a new run')
p.runs # aout put runs

p.runs[1].bold = True

doc.save('/document3.docx')


# Grab the entire text from a docx document
import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for paragraph in doc.paragraphs:
        fullText.append(paragraph.text)
    
    return '\n'.join(fullText)











